from __future__ import annotations

import html
import json
import os
import re
import shutil
import time
import textwrap
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "assets"
OUT_HTML = ROOT / "london_birthday_itinerary.html"
OUT_DOCX = ROOT / "london_birthday_itinerary.docx"
OUT_README = ROOT / "README_london_birthday_itinerary.md"


HOTEL = {
    "name": "Holiday Inn Express London - Victoria",
    "address_lines": [
        "106-110 Belgrave Road",
        "London SW1V 2BJ",
        "United Kingdom",
    ],
    "phone": "+44 20 7630 8888",
}


MAP_QUERIES = {
    "Hotel": "Holiday Inn Express London - Victoria, 106-110 Belgrave Road, London SW1V 2BJ",
    "Tachbrook Street Market": "Tachbrook Street Market, Pimlico",
    "Victoria Station Buckingham Palace Road entrance": "Victoria Station Buckingham Palace Road entrance",
    "London Eye": "London Eye, London",
    "Westminster Bridge": "Westminster Bridge, London",
    "Big Ben": "Big Ben, London",
    "Parliament Square": "Parliament Square, London",
    "Westminster Abbey": "Westminster Abbey, London",
    "Southbank Centre / Royal Festival Hall": "Southbank Centre Royal Festival Hall, London",
    "Gabriel's Wharf": "Gabriel's Wharf, London",
    "Oxo Tower": "Oxo Tower, London",
    "Pimlico Station": "Pimlico Station, Bessborough Street, London SW1V 2JA",
    "Tower Hill Station": "Tower Hill Station, London",
    "Tower of London": "Tower of London",
    "Tower Bridge": "Tower Bridge, London",
    "Borough Market": "Borough Market, London",
    "London Bridge Underground Station": "London Bridge Underground Station, London",
    "Leicester Square Station": "Leicester Square Station, London",
    "Covent Garden": "Covent Garden, London",
    "Seven Dials": "Seven Dials, London",
    "Neal's Yard": "Neal's Yard, Covent Garden, London",
    "Soho": "Soho, London",
    "Carnaby Street": "Carnaby Street, London",
    "Chinatown": "Chinatown Gate, London",
    "Buckingham Palace": "Buckingham Palace, London",
    "St. James's Park": "St James's Park, London",
    "The Mall": "The Mall, London",
    "Trafalgar Square": "Trafalgar Square, London",
    "Camden Market": "Camden Market, London",
    "Camden Town Station": "Camden Town Station, London",
    "Regent's Canal": "Regent's Canal Camden Lock, London",
}


def maps_url(query: str) -> str:
    return "https://www.google.com/maps/search/?api=1&query=" + urllib.parse.quote(query)


PHOTO_QUERIES = [
    ("Hotel area / Pimlico", "Pimlico London Belgrave Road", "pimlico"),
    ("Big Ben", "Big Ben from Westminster Bridge London", "big_ben"),
    ("London Eye", "London Eye Westminster Bridge", "london_eye"),
    ("Westminster Bridge", "Westminster Bridge London", "westminster_bridge"),
    ("South Bank / Queen's Walk", "Queen's Walk South Bank London", "south_bank"),
    ("Tower Bridge", "Tower Bridge from south bank", "tower_bridge"),
    ("Borough Market", "Borough Market London", "borough_market"),
    ("Covent Garden", "Covent Garden London market", "covent_garden"),
    ("Neal's Yard", "Neal's Yard Covent Garden London", "neals_yard"),
    ("Chinatown", "Chinatown Gate London", "chinatown"),
    ("Buckingham Palace", "Buckingham Palace gates London", "buckingham_palace"),
    ("St. James's Park", "St James Park London", "st_jamess_park"),
    ("Camden Market", "Camden Market London", "camden_market"),
]

STATIC_PHOTOS = {
    "london_eye": {
        "title": "The London Eye, from Westminster Bridge.JPG",
        "credit": "Evans1551, public domain",
        "source": "https://commons.wikimedia.org/wiki/File:The_London_Eye,_from_Westminster_Bridge.JPG",
    },
    "westminster_bridge": {
        "title": "Westminster Bridge and Big Ben.jpg",
        "credit": "Ralf Roletschek, Wikimedia Commons license",
        "source": "https://commons.wikimedia.org/wiki/File:Westminster_Bridge_and_Big_Ben.jpg",
    },
    "tower_bridge": {
        "title": "Tower Bridge from South Bank.jpg",
        "credit": "Helen Ilus, Wikimedia Commons license",
        "source": "https://commons.wikimedia.org/wiki/File:Tower_Bridge_from_South_Bank.jpg",
    },
    "borough_market": {
        "title": "Borough Market in London.jpg",
        "credit": "Jeong seolah, CC0",
        "source": "https://commons.wikimedia.org/wiki/File:Borough_Market_in_London.jpg",
    },
    "neals_yard": {
        "title": "Neal's Yard (39823638025).jpg",
        "credit": "Paul The Archivist, CC BY-SA 2.0",
        "source": "https://commons.wikimedia.org/wiki/File:Neal%27s_Yard_(39823638025).jpg",
    },
    "camden_market": {
        "title": "Camden Market - London.jpg",
        "credit": "Jim Linwood, CC BY 2.0",
        "source": "https://commons.wikimedia.org/wiki/File:Camden_Market_-_London.jpg",
    },
}


DAYS = [
    {
        "num": "Day 1",
        "date": "Friday, June 26",
        "title": "Victoria, Westminster and South Bank",
        "area": "Victoria / Westminster / South Bank",
        "transport": "Walking + Big Bus + Uber home",
        "food": "Casual lunch near hotel; casual dinner at Southbank Centre / Royal Festival Hall",
        "tickets": "Big Bus and London Eye",
        "night": "Uber directly back to the hotel",
        "photo": "london_eye",
        "steps": [
            ("Eat near the hotel", "Walk to Tachbrook Street / Warwick Way for an easy cafe or casual restaurant.", ["Tachbrook Street Market"]),
            ("Start the bus loop", "Walk to the Victoria Station / Buckingham Palace Road entrance. Board the Hop-On / Hop-Off Big Bus that matches your ticket.", ["Victoria Station Buckingham Palace Road entrance"]),
            ("Sightseeing loop", "Stay on the bus for the main sightseeing loop. You should pass Buckingham Palace / Green Park, Trafalgar Square / Piccadilly, Westminster / Big Ben, and London Eye / South Bank.", []),
            ("Westminster photos", "Get off near London Eye / Westminster Bridge. Walk: London Eye -> Westminster Bridge -> Big Ben photos -> Parliament Square -> Westminster Abbey exterior -> back to London Eye.", ["London Eye", "Westminster Bridge", "Big Ben", "Parliament Square", "Westminster Abbey"]),
            ("London Eye", "Ride the London Eye. Aim for a timed ticket around 5:30 PM or 6:00 PM.", ["London Eye"]),
            ("Dinner and river walk", "After the London Eye, walk along Queen's Walk. Eat around Southbank Centre / Royal Festival Hall. If you still have energy, continue toward Gabriel's Wharf or Oxo Tower.", ["Southbank Centre / Royal Festival Hall", "Gabriel's Wharf", "Oxo Tower"]),
            ("Go home", "Take an Uber directly back to the hotel. Night 1 return is Uber or black cab, not the Tube.", ["Hotel"]),
        ],
        "photo_spots": ["Big Ben from Westminster Bridge", "London Eye from Westminster Bridge", "Queen's Walk river views"],
        "tired": "If tired after the London Eye, skip the longer Queen's Walk and eat near Southbank Centre / Royal Festival Hall, then Uber directly back to the hotel.",
        "rain": "If raining, still do the London Eye if tickets are booked, then stay around Southbank Centre / Royal Festival Hall for food and indoor cover. Use Uber back to the hotel.",
    },
    {
        "num": "Day 2",
        "date": "Saturday, June 27",
        "title": "Tower Bridge, Borough Market and West End Exploring",
        "area": "Tower Hill / Tower Bridge / Borough Market / Covent Garden / Soho / Chinatown",
        "transport": "Tube + walking; Uber or black cab back if tired",
        "food": "Borough Market lunch; casual dinner in Soho / Chinatown / Covent Garden",
        "tickets": "No tickets planned in advance",
        "night": "Tube if early and comfortable; Uber or black cab if tired",
        "photo": "tower_bridge",
        "steps": [
            ("Breakfast", "Start with breakfast at the hotel.", []),
            ("Tube to Tower Hill", "Walk to Pimlico Station. Take Pimlico -> Victoria Station, change at Victoria, then Victoria -> Tower Hill Station. Confirm the easiest route that morning with TfL Go or Google Maps.", ["Pimlico Station", "Tower Hill Station"]),
            ("Tower area", "Walk around the Tower of London area. No need to buy tickets ahead: see it from the outside, take photos, and enjoy the area without committing to a long tour.", ["Tower of London"]),
            ("Tower Bridge", "Walk to Tower Bridge, take photos outside, and walk across for the views. Only buy Tower Bridge tickets if you decide in the moment that you want to go inside.", ["Tower Bridge"]),
            ("Borough Market lunch", "Walk from Tower Bridge along the south side of the Thames toward London Bridge / Borough Market. Saturday will be crowded, so walk around first and pick food casually.", ["Borough Market", "London Bridge Underground Station"]),
            ("West End exploring", "Take the Northern line northbound from London Bridge to Leicester Square. Walk Leicester Square -> Covent Garden -> Seven Dials -> Neal's Yard -> Soho -> Carnaby Street -> Chinatown.", ["Leicester Square Station", "Covent Garden", "Seven Dials", "Neal's Yard", "Soho", "Carnaby Street", "Chinatown"]),
            ("Dinner and return", "Dinner in Soho, Chinatown, or Covent Garden. Use the Tube if it is still early and comfortable; use Uber or black cab if tired, late, or pickup is easier.", ["Hotel"]),
        ],
        "photo_spots": ["Tower Bridge from the south side of the Thames", "Neal's Yard colorful courtyard", "Chinatown gate", "Carnaby Street"],
        "tired": "If tired after Borough Market, skip the full Soho / Carnaby Street walk and go straight to Covent Garden and Neal's Yard, then dinner.",
        "rain": "If raining, spend more time around Borough Market, Covent Garden, covered shops, cafes, and Chinatown. Use Tube or Uber between areas as needed.",
    },
    {
        "num": "Day 3",
        "date": "Sunday, June 28",
        "title": "Palace Morning and Camden Daytime Adventure",
        "area": "Buckingham Palace / St. James's Park / Trafalgar Square / Camden",
        "transport": "Walking + Tube; Uber or Tube back depending on timing and energy",
        "food": "Camden Market lunch; final dinner in Covent Garden, Soho, or near the hotel",
        "tickets": "No fixed tickets",
        "night": "Keep Camden as a daytime stop; return central before final dinner",
        "photo": "camden_market",
        "steps": [
            ("Breakfast", "Start with breakfast at the hotel.", []),
            ("Palace photos", "Walk or take short transit to Buckingham Palace. See the exterior and gates.", ["Buckingham Palace"]),
            ("Park and Mall walk", "Walk through St. James's Park, then walk The Mall toward Trafalgar Square.", ["St. James's Park", "The Mall", "Trafalgar Square"]),
            ("Tube to Camden", "From central London, take the Tube to Camden Town. Use Google Maps or TfL Go for the best route that morning.", ["Camden Town Station"]),
            ("Camden Market lunch", "Make Camden Market the daytime anchor. Get lunch and explore the stalls, shops, signs, and street food.", ["Camden Market"]),
            ("Regent's Canal if energy is good", "If weather and energy are good, take a short Regent's Canal walk near Camden Lock. Keep it short and easy.", ["Regent's Canal"]),
            ("Final dinner", "Return toward central London for final dinner in Covent Garden, Soho, or near the hotel. Use Uber or the Tube depending on timing and energy.", ["Covent Garden", "Soho", "Hotel"]),
        ],
        "photo_spots": ["Buckingham Palace gates", "Camden Market signs", "Queen's Walk river views"],
        "tired": "If tired, skip Regent's Canal and spend only 1 to 2 hours at Camden Market before returning toward the hotel.",
        "rain": "If raining, shorten the park walk and spend more time in Camden Market covered areas, cafes, and shops.",
    },
]


TODO = [
    "Order British pounds from Chase",
    "Download TripIt",
    "Download JetBlue app",
    "Download offline maps for London",
    "Download TfL Go",
    "Download Uber",
    "Download FREENOW",
    "Apply for UK ETA for Tiffany and Collin",
    "Buy Big Bus London hop-on hop-off tickets",
    "Buy London Eye tickets",
    "Save hotel address as favorite in Uber",
    "Save hotel address in Google Maps",
    "Save parent travel consent letter on both phones",
]

BRING = [
    "Passports",
    "International chargers / UK plug adapters",
    "Copies of UK ETA confirmations",
    "Copy of parental travel consent letter",
    "Printed hotel confirmation",
    "Printed return flight confirmation",
    "Portable phone charger",
    "Comfortable shoes",
    "Rain jacket or small umbrella",
    "Small amount of cash",
    "Credit/debit card",
    "Medication, if applicable",
]

TICKETS = [
    "Big Bus ticket",
    "London Eye ticket",
    "Flight confirmation",
    "Hotel confirmation",
    "UK ETA confirmations",
    "Parent travel consent letter",
]

APPS = ["Google Maps", "TfL Go", "Uber", "FREENOW", "JetBlue app", "TripIt"]

FOOD_AREAS = [
    ("Near hotel", "Tachbrook Street / Warwick Way for easy food near the hotel."),
    ("Day 1 dinner", "Southbank Centre / Royal Festival Hall for casual walk-in options."),
    ("Day 2 lunch", "Borough Market. Walk around first, then choose."),
    ("West End", "Chinatown, Soho, Covent Garden, and Seven Dials for casual dinner, snacks, and dessert."),
    ("Day 3 lunch", "Camden Market for street food and snacks."),
]


def slug_filename(name: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9]+", "_", name.lower()).strip("_")
    return name or "image"


def request_json(url: str) -> dict:
    req = urllib.request.Request(url, headers={"User-Agent": "Codex travel guide builder/1.0"})
    with urllib.request.urlopen(req, timeout=30) as response:
        return json.loads(response.read().decode("utf-8"))


def commons_image(label: str, query: str, slug: str) -> dict:
    if slug in STATIC_PHOTOS:
        static = STATIC_PHOTOS[slug]
        path = ASSET_DIR / f"{slug}.jpg"
        file_url = "https://commons.wikimedia.org/wiki/Special:FilePath/" + urllib.parse.quote(static["title"]) + "?width=1280"
        try:
            req = urllib.request.Request(file_url, headers={"User-Agent": "Codex travel guide builder/1.0"})
            with urllib.request.urlopen(req, timeout=30) as response:
                path.write_bytes(response.read())
            normalize_image(path)
        except Exception:
            if not path.exists():
                return make_placeholder(label, slug)
        return {
            "label": label,
            "slug": slug,
            "path": path,
            "file": path.name,
            "source": static["source"],
            "credit": static["credit"],
            "title": static["title"],
        }
    for existing in [ASSET_DIR / f"{slug}.jpg", ASSET_DIR / f"{slug}.jpeg", ASSET_DIR / f"{slug}.png"]:
        if existing.exists():
            return {
                "label": label,
                "slug": slug,
                "path": existing,
                "file": existing.name,
                "source": "https://commons.wikimedia.org/",
                "credit": "Wikimedia Commons image",
                "title": label,
            }
    search_url = "https://commons.wikimedia.org/w/api.php?" + urllib.parse.urlencode(
        {
            "action": "query",
            "format": "json",
            "generator": "search",
            "gsrsearch": f'filetype:bitmap {query}',
            "gsrnamespace": "6",
            "gsrlimit": "8",
            "prop": "imageinfo",
            "iiprop": "url|extmetadata",
            "iiurlwidth": "1200",
        }
    )
    try:
        data = request_json(search_url)
    except Exception:
        return make_placeholder(label, slug)
    pages = list(data.get("query", {}).get("pages", {}).values())
    pages = [p for p in pages if p.get("imageinfo")]
    if not pages:
        return make_placeholder(label, slug)

    def score(page: dict) -> int:
        title = page.get("title", "").lower()
        bad = ["map", "diagram", "logo", "floor", "signage"]
        return sum(word in title for word in query.lower().split()) - 3 * sum(b in title for b in bad)

    page = sorted(pages, key=score, reverse=True)[0]
    info = page["imageinfo"][0]
    url = info.get("thumburl") or info.get("url")
    ext = Path(urllib.parse.urlparse(url).path).suffix.lower()
    if ext not in [".jpg", ".jpeg", ".png", ".webp"]:
        ext = ".jpg"
    path = ASSET_DIR / f"{slug}{ext}"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Codex travel guide builder/1.0"})
        with urllib.request.urlopen(req, timeout=30) as response:
            path.write_bytes(response.read())
        normalize_image(path)
    except Exception:
        return make_placeholder(label, slug)

    meta = info.get("extmetadata", {})
    artist = clean_meta(meta.get("Artist", {}).get("value", "Wikimedia Commons contributor"))
    license_short = clean_meta(meta.get("LicenseShortName", {}).get("value", "Wikimedia Commons license"))
    desc_url = info.get("descriptionurl", "https://commons.wikimedia.org/")
    return {
        "label": label,
        "slug": slug,
        "path": path,
        "file": path.name,
        "source": desc_url,
        "credit": f"{artist}, {license_short}",
        "title": page.get("title", label).replace("File:", ""),
    }


def clean_meta(value: str) -> str:
    value = re.sub(r"<[^>]+>", "", value or "")
    value = html.unescape(value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:160] or "Wikimedia Commons contributor"


def normalize_image(path: Path) -> None:
    with Image.open(path) as img:
        img = img.convert("RGB")
        img.thumbnail((1600, 1000), Image.LANCZOS)
        if path.suffix.lower() not in [".jpg", ".jpeg"]:
            new_path = path.with_suffix(".jpg")
            img.save(new_path, "JPEG", quality=88)
            path.unlink(missing_ok=True)
            return
        img.save(path, "JPEG", quality=88)


def make_placeholder(label: str, slug: str) -> dict:
    path = ASSET_DIR / f"{slug}.jpg"
    img = Image.new("RGB", (1200, 750), "#e8f2f1")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1200, 750), fill="#e8f2f1")
    draw.rectangle((70, 70, 1130, 680), outline="#86a6a3", width=8)
    text = f"Photo placeholder\n{label}"
    draw.multiline_text((120, 285), text, fill="#294c4a", spacing=18)
    img.save(path, "JPEG", quality=90)
    return {
        "label": label,
        "slug": slug,
        "path": path,
        "file": path.name,
        "source": "",
        "credit": "Placeholder image",
        "title": label,
    }


def build_assets() -> dict[str, dict]:
    ASSET_DIR.mkdir(exist_ok=True)
    photos = {}
    for label, query, slug in PHOTO_QUERIES:
        photos[slug] = commons_image(label, query, slug)
        time.sleep(0.8)
    return photos


def button(label: str, query_name: str) -> str:
    query = MAP_QUERIES[query_name]
    return f'<a class="map-btn" href="{maps_url(query)}" target="_blank" rel="noopener">Open in Google Maps</a>'


def render_map_links(names: list[str]) -> str:
    if not names:
        return ""
    links = []
    for name in names:
        links.append(f'<span class="map-chip"><span>{html.escape(name)}</span>{button("Open", name)}</span>')
    return '<div class="map-links">' + "\n".join(links) + "</div>"


def build_html(photos: dict[str, dict]) -> None:
    hotel_box = "\n".join([HOTEL["name"], *HOTEL["address_lines"], f"Phone: {HOTEL['phone']}"])
    photo_cards = "\n".join(
        f'''
        <figure class="photo-card">
          <img src="assets/{p['file']}" alt="{html.escape(p['label'])}">
          <figcaption>{html.escape(p['label'])}<br><a href="{html.escape(p['source'])}" target="_blank" rel="noopener">{html.escape(p['credit'])}</a></figcaption>
        </figure>'''
        for p in photos.values()
    )
    day_sections = []
    for day in DAYS:
        steps = []
        for i, (title, body, links) in enumerate(day["steps"], 1):
            steps.append(
                f'''
                <li>
                  <div class="step-num">{i}</div>
                  <div>
                    <h4>{html.escape(title)}</h4>
                    <p>{html.escape(body)}</p>
                    {render_map_links(links)}
                  </div>
                </li>'''
            )
        spots = "".join(f"<li>{html.escape(s)}</li>" for s in day["photo_spots"])
        photo = photos[day["photo"]]
        day_sections.append(
            f'''
            <section class="day-card" id="{day['num'].lower().replace(' ', '-')}">
              <img class="day-hero" src="assets/{photo['file']}" alt="{html.escape(photo['label'])}">
              <div class="day-body">
                <p class="eyebrow">{html.escape(day['num'])} - {html.escape(day['date'])}</p>
                <h2>{html.escape(day['title'])}</h2>
                <div class="glance-grid">
                  <div><strong>Main area</strong><span>{html.escape(day['area'])}</span></div>
                  <div><strong>Transportation</strong><span>{html.escape(day['transport'])}</span></div>
                  <div><strong>Food</strong><span>{html.escape(day['food'])}</span></div>
                  <div><strong>Tickets</strong><span>{html.escape(day['tickets'])}</span></div>
                  <div><strong>Night return</strong><span>{html.escape(day['night'])}</span></div>
                </div>
                <ol class="steps">{''.join(steps)}</ol>
                <div class="callout-row">
                  <div class="callout photo"><h3>Photo Spots</h3><ul>{spots}</ul></div>
                  <div class="callout tired"><h3>If You're Tired</h3><p>{html.escape(day['tired'])}</p></div>
                  <div class="callout rain"><h3>Rain Plan</h3><p>{html.escape(day['rain'])}</p></div>
                </div>
              </div>
            </section>'''
        )

    maps_list = "\n".join(
        f'<a href="{maps_url(query)}" target="_blank" rel="noopener"><span>{html.escape(name)}</span><b>Open</b></a>'
        for name, query in MAP_QUERIES.items()
    )
    todo_items = "".join(f'<label><input type="checkbox"> {html.escape(item)}</label>' for item in TODO)
    bring_items = "".join(f'<label><input type="checkbox"> {html.escape(item)}</label>' for item in BRING)
    ticket_items = "".join(f'<div class="ticket-slot"><strong>{html.escape(t)}</strong><span>Add confirmation number, QR code note, or screenshot location.</span></div>' for t in TICKETS)
    apps = "".join(f"<li>{html.escape(app)}</li>" for app in APPS)
    food = "".join(f"<li><strong>{html.escape(k)}:</strong> {html.escape(v)}</li>" for k, v in FOOD_AREAS)

    css = r"""
    :root{--ink:#102828;--muted:#5d6d6d;--sea:#dcefed;--mint:#edf8f2;--coral:#ff7f6e;--gold:#f7c56b;--sky:#dcecff;--line:#d8e3e1;--card:#ffffff}
    *{box-sizing:border-box} body{margin:0;font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",Arial,sans-serif;color:var(--ink);background:#f7faf9;line-height:1.55}
    a{color:#12605b} .hero{min-height:74vh;background:linear-gradient(180deg,rgba(12,34,38,.40),rgba(12,34,38,.70)),url("assets/london_eye.jpg") center/cover;display:flex;align-items:end;padding:24px}
    .hero-inner{max-width:1000px;margin:0 auto;color:white;width:100%;padding-bottom:34px}.kicker{font-size:.8rem;text-transform:uppercase;letter-spacing:.12em;font-weight:800}.hero h1{font-size:clamp(2.2rem,7vw,5rem);line-height:.95;margin:.3rem 0}.hero p{font-size:1.1rem;max-width:760px}
    .topnav{position:sticky;top:0;z-index:2;background:rgba(255,255,255,.94);backdrop-filter:blur(12px);border-bottom:1px solid var(--line);display:flex;gap:8px;overflow-x:auto;padding:10px 16px}.topnav a{white-space:nowrap;text-decoration:none;font-weight:750;padding:9px 12px;border-radius:999px;background:var(--mint);color:#173f3d}
    main{max-width:1120px;margin:0 auto;padding:18px}.section{margin:24px 0}.section h2,.day-body h2{font-size:clamp(1.55rem,4vw,2.5rem);line-height:1.05;margin:0 0 14px}.grid{display:grid;gap:14px}.two{grid-template-columns:repeat(2,minmax(0,1fr))}.three{grid-template-columns:repeat(3,minmax(0,1fr))}
    .card,.quick,.ticket-slot,.callout{background:var(--card);border:1px solid var(--line);border-radius:8px;padding:16px;box-shadow:0 10px 28px rgba(22,51,48,.07)}.quick{background:linear-gradient(135deg,#ffffff,#edf8f2)}
    .copybox{font-family:ui-monospace,SFMono-Regular,Consolas,monospace;background:#102828;color:white;padding:14px;border-radius:8px;white-space:pre-wrap;margin:12px 0}.copy-btn,.map-btn{display:inline-flex;align-items:center;justify-content:center;border:0;border-radius:8px;background:#173f3d;color:white;text-decoration:none;font-weight:800;padding:9px 12px;cursor:pointer}
    .checklist{display:grid;gap:8px}.checklist label{background:white;border:1px solid var(--line);border-radius:8px;padding:10px}.apps li{margin:6px 0}.day-card{background:white;border:1px solid var(--line);border-radius:8px;overflow:hidden;margin:26px 0;box-shadow:0 16px 38px rgba(22,51,48,.08)}.day-hero{width:100%;height:270px;object-fit:cover}.day-body{padding:18px}.eyebrow{font-weight:900;color:#b9473b;text-transform:uppercase;font-size:.78rem;letter-spacing:.09em}
    .glance-grid{display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:10px;margin:14px 0}.glance-grid div{background:#f5faf9;border:1px solid var(--line);border-radius:8px;padding:11px}.glance-grid strong{display:block;font-size:.75rem;text-transform:uppercase;color:var(--muted)}.glance-grid span{display:block;font-weight:700;font-size:.93rem}
    .steps{list-style:none;padding:0;margin:18px 0;display:grid;gap:14px}.steps li{display:grid;grid-template-columns:42px 1fr;gap:12px}.step-num{width:34px;height:34px;border-radius:50%;background:var(--coral);color:white;display:grid;place-items:center;font-weight:900}.steps h4{margin:0 0 3px;font-size:1.05rem}.steps p{margin:0 0 8px}
    .map-links{display:flex;flex-wrap:wrap;gap:8px}.map-chip{display:flex;gap:7px;align-items:center;border:1px solid var(--line);border-radius:8px;padding:6px;background:#fbfdfd}.map-chip span{font-size:.88rem}.map-chip .map-btn{font-size:.78rem;padding:6px 8px}.callout-row{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:12px}.callout h3{margin:0 0 7px}.photo{background:#fff8e8}.tired{background:#eef4ff}.rain{background:#eef8f2}
    .photo-grid{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:12px}.photo-card{margin:0;background:white;border:1px solid var(--line);border-radius:8px;overflow:hidden}.photo-card img{width:100%;height:170px;object-fit:cover}.photo-card figcaption{font-size:.78rem;color:var(--muted);padding:9px}.map-index{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:8px}.map-index a{display:flex;justify-content:space-between;gap:10px;padding:10px;border:1px solid var(--line);background:white;border-radius:8px;text-decoration:none}.transport p{margin:.35rem 0}.footer{color:var(--muted);font-size:.85rem;margin:28px 0}
    @media(max-width:820px){.two,.three,.glance-grid,.callout-row,.photo-grid,.map-index{grid-template-columns:1fr}.hero{min-height:64vh;padding:18px}.day-hero{height:210px}main{padding:12px}.map-chip{width:100%;justify-content:space-between}.steps li{grid-template-columns:36px 1fr}.topnav{padding:9px 10px}.quick{padding:14px}}
    @media print{.topnav,.copy-btn,.map-btn{display:none}.hero{min-height:auto;color:#102828;background:none;padding:0}.hero-inner{color:#102828;padding:0}.card,.quick,.day-card,.ticket-slot,.callout{box-shadow:none}.day-card{break-inside:avoid}.photo-grid{grid-template-columns:repeat(2,1fr)}body{background:white}main{max-width:none}.day-hero{height:180px}}
    """
    script = """
    function copyHotel(){
      const text = document.getElementById('hotel-copy').innerText;
      navigator.clipboard.writeText(text);
      const btn = document.getElementById('copy-btn');
      btn.innerText = 'Copied';
      setTimeout(()=>btn.innerText='Copy hotel address', 1500);
    }
    """
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>London Birthday Itinerary - Tiffany and Collin</title>
  <style>{css}</style>
</head>
<body>
  <header class="hero">
    <div class="hero-inner">
      <div class="kicker">June 26-29, 2026</div>
      <h1>London Birthday Trip</h1>
      <p>A first-time London guide for Tiffany Bediner and Collin Bediner: practical, flexible, photo-friendly, and easy to follow from a phone.</p>
    </div>
  </header>
  <nav class="topnav">
    <a href="#quick">Quick Reference</a><a href="#todo">To Do</a><a href="#tickets">Ticket Wallet</a><a href="#day-1">Day 1</a><a href="#day-2">Day 2</a><a href="#day-3">Day 3</a><a href="#maps">Maps</a>
  </nav>
  <main>
    <section class="section quick" id="quick">
      <h2>Quick Reference</h2>
      <div class="grid two">
        <div>
          <p><strong>Travelers:</strong> Tiffany Bediner, 20; Collin Bediner, 14<br>
          <strong>Trip dates:</strong> Travel starts Thursday, June 25, 2026; arrive London Friday morning, June 26; depart London Monday, June 29.</p>
          <p><strong>Hotel:</strong> {HOTEL['name']}<br>
          <strong>Address:</strong> {', '.join(HOTEL['address_lines'])}<br>
          <strong>Phone:</strong> {HOTEL['phone']}<br>
          <strong>Nearest Tube:</strong> Pimlico<br>
          <strong>Main nearby hub:</strong> Victoria Station<br>
          <strong>UK emergency number:</strong> 999</p>
          {button('Open in Google Maps', 'Hotel')}
        </div>
        <div>
          <p><strong>Parent contacts:</strong><br>Parent 1: ____________________<br>Parent 2: ____________________</p>
          <p>Save the hotel in Uber and Google Maps before leaving home. Use Uber or FREENOW when tired; use only Uber, FREENOW, or an official black cab.</p>
        </div>
      </div>
      <h3>Copy / Paste Hotel Address</h3>
      <div class="copybox" id="hotel-copy">{html.escape(hotel_box)}</div>
      <button class="copy-btn" id="copy-btn" onclick="copyHotel()">Copy hotel address</button>
    </section>

    <section class="section grid two" id="todo">
      <div class="card"><h2>To Do Before Travel</h2><div class="checklist">{todo_items}</div></div>
      <div class="card"><h2>To Bring</h2><div class="checklist">{bring_items}</div></div>
    </section>

    <section class="section card" id="tickets"><h2>Ticket Wallet</h2><p>Keep screenshots saved on both phones, plus printed copies for the most important items.</p><div class="grid three">{ticket_items}</div></section>

    <section class="section card transport">
      <h2>Transportation Tips</h2>
      <p><strong>Use the Tube for most daytime travel.</strong> Each person needs their own contactless card, phone, or Oyster card. Use the same device to tap in and tap out. Do not tap in with a phone and tap out with an Apple Watch or a different card.</p>
      <p><strong>Keep phone battery charged.</strong> Use TfL Go and Google Maps to confirm routes each morning.</p>
      <p><strong>Use Uber or a black cab at night if tired.</strong> Uber is usually cheaper. A black cab may feel easier and more official if one is available. FREENOW can book black cabs.</p>
      <p><strong>Never accept rides from anyone offering a car on the street.</strong> Use only Uber, FREENOW, or an official black cab with a TAXI light.</p>
      <p><strong>Night 1 rule:</strong> Do not take the Tube back on Night 1. Take Uber directly to the hotel. If Uber is too expensive, confusing, or pickup looks difficult, use FREENOW or an official black cab.</p>
    </section>

    <section class="section grid two">
      <div class="card"><h2>Apps To Have</h2><ul class="apps">{apps}</ul></div>
      <div class="card"><h2>Safety and Common Sense</h2><p>Stay together. Stay on main streets and main walking routes. Keep phones zipped away when not actively using them. Use Uber or black cab if tired or if it feels late. Do not accept rides from strangers. Message parents when leaving the hotel, arriving at major stops, and heading back. Carry passport copies, not passports, during daily sightseeing unless otherwise needed. Use the hotel address if you ever need help getting back.</p></div>
    </section>

    <section class="section card"><h2>Food Strategy</h2><p>Do not over-plan restaurants. Most meals should be casual walk-in areas so the day can flex around weather, energy, and crowds.</p><ul>{food}</ul></section>

    {''.join(day_sections)}

    <section class="section card" id="departure">
      <h2>Departure Day: Monday, June 29</h2>
      <p>Keep this morning simple. Eat breakfast at the hotel if timing works, double-check passports and phones before leaving the room, and use the JetBlue app plus the flight confirmation for updates. Use the hotel address and saved ride apps if you need help coordinating pickup.</p>
    </section>

    <section class="section" id="photos"><h2>Photo Gallery and Sources</h2><div class="photo-grid">{photo_cards}</div></section>
    <section class="section card" id="maps"><h2>Map Link Index</h2><div class="map-index">{maps_list}</div></section>
    <p class="footer">Created for Tiffany Bediner and Collin Bediner. All attraction plans stay inside London.</p>
  </main>
  <script>{script}</script>
</body>
</html>"""
    OUT_HTML.write_text(html_doc, encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, bold=False, color="102828"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", style=None, size=10.5, bold=False, color="102828", after=5, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.12
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    return p


def h(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run(r, size={1: 17, 2: 13.5, 3: 11.5}[level], bold=True, color={1: "173F3D", 2: "B9473B", 3: "102828"}[level])
    return p


def bullet(doc, text, indent=0.25):
    p = para(doc, "", after=3)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run("- ")
    set_run(r, bold=True)
    r = p.add_run(text)
    set_run(r)
    return p


def add_callout(doc, title, text, fill="EDF8F2"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, 10.5, True, "173F3D")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r = p2.add_run(text)
    set_run(r, 10)
    para(doc, "", after=4)


def add_key_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_margins(c0)
        set_cell_margins(c1)
        set_cell_shading(c0, "EDF8F2")
        p = c0.paragraphs[0]
        r = p.add_run(label)
        set_run(r, 9, True, "173F3D")
        p = c1.paragraphs[0]
        r = p.add_run(value)
        set_run(r, 9.5)
    para(doc, "", after=4)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "12605B")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx(photos: dict[str, dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("London Birthday Trip")
    set_run(r, 25, True, "173F3D")
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tiffany Bediner and Collin Bediner | June 26-29, 2026")
    set_run(r, 11, False, "5D6D6D")
    p.paragraph_format.space_after = Pt(12)

    try:
        doc.add_picture(str(photos["london_eye"]["path"]), width=Inches(7.0))
    except Exception:
        pass

    h(doc, "Quick Reference", 1)
    para(doc, "Travelers: Tiffany Bediner, 20; Collin Bediner, 14", bold=True)
    para(doc, "Trip dates: Travel starts Thursday, June 25, 2026; arrive London Friday morning, June 26; depart London Monday, June 29.", bold=True)
    add_key_table(
        doc,
        [
            ("Hotel", HOTEL["name"]),
            ("Address", ", ".join(HOTEL["address_lines"])),
            ("Phone", HOTEL["phone"]),
            ("Nearest Tube", "Pimlico"),
            ("Nearby hub", "Victoria Station"),
            ("Emergency", "999"),
            ("Parent contacts", "Parent 1: ____________________   Parent 2: ____________________"),
        ],
    )
    add_callout(doc, "Copy / Paste Hotel Address", "\n".join([HOTEL["name"], *HOTEL["address_lines"], f"Phone: {HOTEL['phone']}"]), "102828")
    para(doc, "Save the hotel in Uber and Google Maps before leaving home. Use Uber or FREENOW when tired; use only Uber, FREENOW, or an official black cab.", bold=True, color="B9473B")

    h(doc, "To Do Before Travel", 1)
    for item in TODO:
        bullet(doc, f"[ ] {item}")
    h(doc, "To Bring", 1)
    for item in BRING:
        bullet(doc, f"[ ] {item}")

    h(doc, "Ticket Wallet", 1)
    for ticket in TICKETS:
        bullet(doc, f"{ticket}: ______________________________________________")

    h(doc, "Transportation Tips", 1)
    for text in [
        "Use the Tube for most daytime travel.",
        "Each person needs their own contactless card, phone, or Oyster card.",
        "Use the same device to tap in and tap out. Do not tap in with a phone and tap out with an Apple Watch or a different card.",
        "Keep phone battery charged. Use TfL Go and Google Maps.",
        "Uber is usually cheaper. A black cab may feel easier and more official if available. FREENOW can book black cabs.",
        "Never accept rides from anyone offering a car on the street. Use only Uber, FREENOW, or an official black cab with a TAXI light.",
        "Night 1 rule: Do not take the Tube back. Take Uber directly back to the hotel. If Uber is too expensive, confusing, or pickup looks difficult, use FREENOW or an official black cab.",
    ]:
        bullet(doc, text)

    h(doc, "Safety and Common Sense", 1)
    para(doc, "Stay together. Stay on main streets and main walking routes. Keep phones zipped away when not actively using them. Use Uber or black cab if tired or if it feels late. Do not accept rides from strangers. Save the hotel in Uber and Google Maps. Keep the portable charger charged. Message parents when leaving the hotel, arriving at major stops, and heading back. Carry passport copies, not passports, during daily sightseeing unless otherwise needed. Use the hotel address if you ever need help getting back.")

    h(doc, "Food Strategy", 1)
    para(doc, "Do not over-plan restaurants. Most meals should be casual walk-in areas instead of reservations.")
    for label, text in FOOD_AREAS:
        bullet(doc, f"{label}: {text}")

    for day in DAYS:
        doc.add_section(WD_SECTION.NEW_PAGE)
        h(doc, f"{day['num']}: {day['date']}", 1)
        para(doc, day["title"], size=12, bold=True, color="B9473B")
        try:
            doc.add_picture(str(photos[day["photo"]]["path"]), width=Inches(6.9))
        except Exception:
            pass
        add_key_table(
            doc,
            [
                ("Main area", day["area"]),
                ("Transportation", day["transport"]),
                ("Food", day["food"]),
                ("Tickets", day["tickets"]),
                ("Night return", day["night"]),
            ],
        )
        h(doc, "Itinerary", 2)
        for i, (title, body, links) in enumerate(day["steps"], 1):
            p = para(doc, f"{i}. {title}: ", bold=True, after=2)
            r = p.add_run(body)
            set_run(r, 10)
            for link_name in links:
                p = para(doc, "   ")
                add_hyperlink(p, maps_url(MAP_QUERIES[link_name]), f"{link_name} - Open in Google Maps")
        h(doc, "Photo Spots", 2)
        for spot in day["photo_spots"]:
            bullet(doc, spot)
        add_callout(doc, "If You're Tired", day["tired"], "EEF4FF")
        add_callout(doc, "Rain Plan", day["rain"], "EEF8F2")

    doc.add_section(WD_SECTION.NEW_PAGE)
    h(doc, "Departure Day: Monday, June 29", 1)
    para(doc, "Keep this morning simple. Eat breakfast at the hotel if timing works, double-check passports and phones before leaving the room, and use the JetBlue app plus the flight confirmation for updates. Use the hotel address and saved ride apps if you need help coordinating pickup.")

    h(doc, "Apps To Download", 1)
    for app in APPS:
        bullet(doc, app)
    h(doc, "Map Link Index", 1)
    for name, query in MAP_QUERIES.items():
        p = para(doc, "")
        r = p.add_run(f"{name}: ")
        set_run(r, 9.5, True)
        add_hyperlink(p, maps_url(query), "Open in Google Maps")

    h(doc, "Photo Sources", 1)
    for pinfo in photos.values():
        p = para(doc, f"{pinfo['label']}: {pinfo['credit']} ")
        if pinfo["source"]:
            add_hyperlink(p, pinfo["source"], "source")

    doc.core_properties.title = "London Birthday Itinerary"
    doc.core_properties.subject = "Trip guide for Tiffany Bediner and Collin Bediner"
    doc.core_properties.author = "Codex"
    doc.save(OUT_DOCX)


def build_readme(photos: dict[str, dict]) -> None:
    lines = [
        "# London Birthday Itinerary Files",
        "",
        "Open `london_birthday_itinerary.html` in a browser for the phone-friendly interactive guide.",
        "Open `london_birthday_itinerary.docx` in Word or Google Docs for the printable version.",
        "If `london_birthday_itinerary.pdf` is present, it was exported from the Word version.",
        "",
        "## Updating Links",
        "",
        "The Google Maps links are generated from place names in `build_london_itinerary.py`. Edit the `MAP_QUERIES` list and rerun the builder if you want a more specific map target.",
        "",
        "## Updating Photos",
        "",
        "Photos are stored in the `assets` folder. Replace any image with a new file using the same filename, or edit `PHOTO_QUERIES` in `build_london_itinerary.py` and rerun the builder.",
        "",
        "## Photo Sources",
        "",
    ]
    for p in photos.values():
        source = p["source"] or "local placeholder"
        lines.append(f"- {p['label']}: {p['credit']} - {source}")
    OUT_README.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    photos = build_assets()
    build_html(photos)
    build_docx(photos)
    build_readme(photos)
    print(f"Wrote {OUT_HTML.name}")
    print(f"Wrote {OUT_DOCX.name}")
    print(f"Wrote {OUT_README.name}")


if __name__ == "__main__":
    main()
